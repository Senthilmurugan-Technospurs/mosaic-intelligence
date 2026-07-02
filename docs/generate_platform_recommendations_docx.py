"""Generate Platform Recommendations Integration spec (Word) for backend team."""
from pathlib import Path
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

OUT = Path(__file__).parent / "MOSAIC-Platform-Recommendations-Integration-Spec.docx"


def heading(doc, text, level=1):
    p = doc.add_heading(text, level=level)
    return p


def para(doc, text, bold=False):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.font.size = Pt(11)
    if bold:
        run.bold = True
    return p


def bullet(doc, text):
    p = doc.add_paragraph(text, style="List Bullet")
    for run in p.runs:
        run.font.size = Pt(11)
    return p


def table(doc, headers, rows):
    t = doc.add_table(rows=1 + len(rows), cols=len(headers))
    t.style = "Table Grid"
    hdr = t.rows[0].cells
    for i, h in enumerate(headers):
        hdr[i].text = h
        for p in hdr[i].paragraphs:
            for r in p.runs:
                r.bold = True
                r.font.size = Pt(10)
    for ri, row in enumerate(rows):
        for ci, val in enumerate(row):
            t.rows[ri + 1].cells[ci].text = str(val)
            for p in t.rows[ri + 1].cells[ci].paragraphs:
                for r in p.runs:
                    r.font.size = Pt(10)
    doc.add_paragraph()
    return t


def code_block(doc, text):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.font.name = "Consolas"
    run.font.size = Pt(9)
    p.paragraph_format.left_indent = Inches(0.25)
    return p


def build():
    doc = Document()
    doc.core_properties.title = "MOSAIC Platform Recommendations — Integration Spec"
    doc.core_properties.author = "MOSAIC Frontend Team"

    # Title
    t = doc.add_heading("MOSAIC Platform Recommendations", 0)
    t.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sub = doc.add_paragraph("Integration specification for backend team — Google Ads & Meta native recommendations")
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for r in sub.runs:
        r.font.size = Pt(12)
        r.font.color.rgb = RGBColor(95, 115, 135)
    doc.add_paragraph("Version: July 2026  |  Status: Demo UI complete; live platform sync not yet implemented")
    doc.add_paragraph()

    # 1. Scope
    heading(doc, "1. Scope", 1)
    para(doc, "This document covers only native ad-platform recommendations (Google Ads and Meta Opportunity Score / Recommendations). It does not cover campaign sync, metrics, insights, or MOSAIC AI recommendation generation.")
    bullet(doc, "Current state: mosaic-frontend uses mock platform recommendations (platformRecommendations.ts).")
    bullet(doc, "Target state: backend sync job pulls live recs, maps to Recommendation type, merges with MOSAIC recs in API responses.")
    bullet(doc, "UI behaviour: Google/Meta recs are view-only in mosaic (Open in Ads Manager). Apply/dismiss for platform recs happens in native managers unless a future phase adds platform apply APIs.")

    # 2. Official APIs
    heading(doc, "2. Official API endpoints (recommendations only)", 1)

    heading(doc, "2.1 Google Ads API", 2)
    para(doc, "Base URL (match existing backend): https://googleads.googleapis.com/v{version}  (default v20 via GOOGLE_ADS_API_VERSION)")
    table(doc, ["Purpose", "Method & endpoint"], [
        ["List recommendations", "POST /customers/{customerId}/googleAds:search"],
        ["List (stream)", "POST /customers/{customerId}/googleAds:searchStream"],
        ["Apply", "POST /customers/{customerId}/recommendations:apply"],
        ["Dismiss", "POST /customers/{customerId}/recommendations:dismiss"],
        ["Generate (optional)", "POST /customers/{customerId}/recommendations:generate"],
    ])
    para(doc, "Documentation: https://developers.google.com/google-ads/api/docs/recommendations")
    para(doc, "GAQL resource: recommendation")
    para(doc, "Required headers (same as existing GoogleAdsService): developer-token, login-customer-id, Authorization bearer token.")
    para(doc, "Example GAQL query:", bold=True)
    code_block(doc, """SELECT
  recommendation.resource_name,
  recommendation.type,
  recommendation.campaign,
  recommendation.ad_group,
  recommendation.status,
  recommendation.impact,
  recommendation.campaign_budget_recommendation,
  recommendation.keyword_recommendation,
  recommendation.target_cpa_opt_in_recommendation
FROM recommendation
WHERE recommendation.campaign = 'customers/{customerId}/campaigns/{googleCampaignId}'
  AND recommendation.dismissed = FALSE""")

    heading(doc, "2.2 Meta Marketing API", 2)
    para(doc, "Base URL: https://graph.facebook.com/v{version}  (backend currently v20.0; Meta docs use v25.0 — align on go-live)")
    table(doc, ["Purpose", "Method & endpoint"], [
        ["List (ad account)", "GET /v{version}/act_{adAccountId}/recommendations"],
        ["List (business portfolio)", "GET /v{version}/{businessId}/recommendations"],
        ["Apply", "POST /v{version}/act_{adAccountId}/recommendations"],
    ])
    para(doc, "Documentation: https://developers.facebook.com/docs/marketing-api/overview/performance-recommendations/")
    para(doc, "Apply requires recommendation_signature and type-specific extra_data from the recommendation object.")
    para(doc, "Example fetch:", bold=True)
    code_block(doc, "GET /v25.0/act_{AD_ACCOUNT_ID}/recommendations?access_token={TOKEN}")

    # 3. Target type
    heading(doc, "3. Target data model — Recommendation", 1)
    para(doc, "All mapped platform recs must conform to the mosaic Recommendation interface (TypeScript). Key fields:")
    table(doc, ["Field", "Platform rule"], [
        ["source", "google_ads | meta"],
        ["isPrimary", "Always false (MOSAIC owns primary)"],
        ["canUndo", "Always false for platform recs in UI"],
        ["generationRunId", "platform-sync or sync batch UUID"],
        ["dataWindow", "Default 7d unless platform provides window"],
        ["level", "campaign (default) or ad_set when ad group/ad set present"],
        ["id", "Stable: google:{resource_name} or meta:{recommendation_signature}"],
        ["externalId", "Platform native ID for dedup/upsert"],
        ["managerDeepLink", "Deep link to native recommendations UI"],
        ["campaignId", "Internal mosaic campaign ID (lookup required)"],
    ])

    # 4. Google mapper
    heading(doc, "4. Google Ads → Recommendation mapper", 1)
    heading(doc, "4.1 Type mapping (recommendation.type → RecType)", 2)
    table(doc, ["Google type(s)", "RecType"], [
        ["CAMPAIGN_BUDGET, FORECASTING_CAMPAIGN_BUDGET, MARGINAL_ROI_CAMPAIGN_BUDGET, MOVE_UNUSED_BUDGET", "budget_orchestration"],
        ["TARGET_CPA_OPT_IN, TARGET_ROAS_OPT_IN, MAXIMIZE_*_OPT_IN, SET_TARGET_CPA, SET_TARGET_ROAS, LOWER_TARGET_ROAS", "bid_policy"],
        ["KEYWORD, KEYWORD_MATCH_TYPE, USE_BROAD_MATCH_KEYWORD", "audience_strategy"],
        ["RESPONSIVE_SEARCH_AD*, IMPROVE_GOOGLE_ADS_*, SITELINK_ASSET, CALLOUT_ASSET, CALL_EXTENSION", "placement"],
        ["Unknown types", "placement (fallback) + log unmapped type"],
    ])
    heading(doc, "4.2 Status mapping", 2)
    table(doc, ["Google", "RecStatus"], [
        ["dismissed = false, not applied", "pending"],
        ["dismissed = true", "dismissed"],
        ["Applied via API", "applied"],
    ])
    heading(doc, "4.3 Score mapping (derived)", 2)
    para(doc, "Google does not provide MOSAIC-style scores. Backend must derive RecScore:")
    bullet(doc, "expectedLift: from recommendation.impact metrics or heuristic (0–100)")
    bullet(doc, "confidence: default 70–80 unless impact data supports higher")
    bullet(doc, "risk: low for budget; medium for bid; high for structural migrations")
    bullet(doc, "reversibility: easy for budget/creative; moderate for bid; hard for migration types")

    # 5. Meta mapper
    heading(doc, "5. Meta → Recommendation mapper", 1)
    heading(doc, "5.1 Type mapping (Meta recommendation type → RecType)", 2)
    table(doc, ["Meta type (examples)", "RecType"], [
        ["BUDGET_LIMITED, INCREASE_BUDGET, budget-related", "budget_orchestration"],
        ["BID_STRATEGY, COST_CAP, bidding-related", "bid_policy"],
        ["AUDIENCE, LOOKALIKE, EXPAND_AUDIENCE", "audience_strategy"],
        ["CREATIVE_FATIGUE, AD_CREATIVE, creative/placement", "placement"],
        ["SCHEDULE, dayparting", "geo_daypart"],
        ["Unknown", "placement (fallback) + log"],
    ])
    heading(doc, "5.2 Key fields", 2)
    bullet(doc, "externalId = recommendation_signature (required for future apply)")
    bullet(doc, "title/description from Meta title, body, message fields")
    bullet(doc, "currentValue/proposedValue parsed from recommendation content (budget %, frequency, audience size, etc.)")

    # 6. Backend service
    heading(doc, "6. Suggested backend service contract", 1)
    code_block(doc, """interface PlatformRecommendationMapper {
  mapGoogle(row, ctx: { internalCampaignId, googleCustomerId, googleCampaignId, syncedAt }): Recommendation;
  mapMeta(item, ctx: { internalCampaignId, metaAdAccountId, syncedAt }): Recommendation;
}

Sync job:
1. For each synced campaign with google_campaign_id / meta_campaign_id
2. Call recommendation endpoints (section 2)
3. Map → Recommendation[]
4. Upsert by (source, externalId)
5. Merge with MOSAIC recs in GET /recommendations/campaigns/{id}""")

    # 7. UI behaviour - platform only groups
    heading(doc, "7. UI behaviour — platform-only recommendations (no MOSAIC equivalent)", 1)
    para(doc, "Yes. If a recommendation category exists on Google or Meta but MOSAIC has nothing in that module, the UI still shows it.", bold=True)
    bullet(doc, "Campaign detail → Google Ads or Meta source tab: platform recs appear in a flat list (no MOSAIC hierarchy).")
    bullet(doc, "Campaign detail → All tab: platform recs appear alongside MOSAIC recs.")
    bullet(doc, "Compare tab / Compare workspace: creates a Platform only row — MOSAIC column empty (—), Google or Meta column populated. Badge: Platform only.")
    bullet(doc, "Module filter: platform recs appear when their mapped RecType matches the filter, or always under All Recommendations.")
    bullet(doc, "Mapper must assign every platform rec to one of the five RecTypes so module filters and compare pairing work. Unmapped Google/Meta types use fallback type + logging.")
    para(doc, "Compare pairing logic groups by rec.type (budget_orchestration, bid_policy, etc.). One platform rec per category per source is paired per row. Platform-only rows are expected and correct.")

    # 8. Go-live checklist
    heading(doc, "8. Go-live checklist (backend)", 1)
    heading(doc, "8.1 Prerequisites", 2)
    bullet(doc, "OAuth credentials per client (Google developer token, refresh token, MCC login-customer-id; Meta access token, act_ account ID)")
    bullet(doc, "Campaign ID bridge table: internal campaignId ↔ google_campaign_id ↔ meta_campaign_id")
    bullet(doc, "API version decision: Google v20+; Meta v20 → v25 upgrade plan")
    bullet(doc, "Rate limit / quota strategy (batch per account, cache TTL e.g. 1–6 hours)")

    heading(doc, "8.2 Implementation deliverables", 2)
    bullet(doc, "platform_recommendations_sync job (scheduled + on-demand)")
    bullet(doc, "GoogleRecommendationMapper + MetaRecommendationMapper")
    bullet(doc, "Upsert store: (source, externalId) unique key")
    bullet(doc, "API: merge platform recs into existing campaign recommendations endpoint")
    bullet(doc, "API: platformSyncedAt timestamp for UI footer")
    bullet(doc, "Logging for unmapped recommendation types")
    bullet(doc, "Error handling: partial sync per campaign (don't fail entire batch)")

    heading(doc, "8.3 Mapper quality requirements", 2)
    bullet(doc, "Populate title, description, currentValue, proposedValue, estimatedImpact — compare alignment uses these strings")
    bullet(doc, "Stable id and externalId for deduplication across sync runs")
    bullet(doc, "managerDeepLink per campaign/account where possible")
    bullet(doc, "Phase 1 type coverage: budget, bid, keyword/audience, creative/RSA")
    bullet(doc, "Phase 2: remaining Google types + Meta edge cases")

    heading(doc, "8.4 Known difficulties", 2)
    table(doc, ["#", "Challenge", "Mitigation"], [
        ["1", "50+ Google recommendation types with different nested payloads", "Phase 1 subset; fallback mapper; log unmapped"],
        ["2", "Meta types vary; need real API samples per client", "Capture samples in staging; defensive parsing"],
        ["3", "No native MOSAIC-style scores", "Derived RecScore with documented heuristics"],
        ["4", "Compare alignment is heuristic (text + category)", "Consistent currentValue/proposedValue formatting"],
        ["5", "Platform recs view-only in UI", "Deep links; optional phase 2 for apply APIs"],
        ["6", "MOSAIC hierarchical vs platform flat", "Platform recs never use parentRecommendationId / isPrimary"],
        ["7", "API version drift", "Pin versions; test before upgrade"],
    ])

    # 9. API response shape
    heading(doc, "9. Expected API response (campaign detail)", 1)
    code_block(doc, """{
  "campaign": { ... },
  "recommendations": [
    { "source": "mosaic", "isPrimary": true, ... },
    { "source": "google_ads", "isPrimary": false, "externalId": "customers/123/recommendations/456", ... },
    { "source": "meta", "isPrimary": false, "externalId": "meta_sig_abc", ... }
  ],
  "comparePairs": [ ... built server-side or client-side from merged list ... ],
  "platformSyncedAt": "2026-07-02T08:00:00Z"
}""")

    # 10. References
    heading(doc, "10. References", 1)
    bullet(doc, "Google Ads Recommendations: https://developers.google.com/google-ads/api/docs/recommendations")
    bullet(doc, "Google recommendation GAQL fields: https://developers.google.com/google-ads/api/fields/v20/recommendation")
    bullet(doc, "Meta Performance Recommendations: https://developers.facebook.com/docs/marketing-api/overview/performance-recommendations/")
    bullet(doc, "MOSAIC frontend types: src/types/recommendation.ts")
    bullet(doc, "Compare logic: src/lib/recommendationCompare.ts")
    bullet(doc, "Demo mocks: src/mocks/platformRecommendations.ts")

    doc.save(OUT)
    print(f"Wrote {OUT}")


if __name__ == "__main__":
    build()
